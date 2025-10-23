from fastapi import UploadFile, File, Form, APIRouter, FastAPI
from fastapi.responses import JSONResponse
import os, base64, uuid
from pdf2image import convert_from_path
from PIL import Image
import pytesseract
from .model import get_client
from .prompts import get_supervision_prompt

# LangChain imports
from langchain.memory import ConversationBufferMemory
from langchain.schema import SystemMessage, HumanMessage, AIMessage


TEMP_FOLDER = os.path.join(os.path.dirname(__file__), "pdf_images")
os.makedirs(TEMP_FOLDER, exist_ok=True)

# Store session memories
session_memories = {}  # { session_id: ConversationBufferMemory }

client = get_client()
DEPLOYMENT_NAME = os.getenv("AZURE_OPENAI_DEPLOYMENT")

assign_route = APIRouter(
    prefix="/assignment",
    tags=["assignment"]
)


# ----------------- START SESSION -----------------
@assign_route.post("/start-session")
async def start_session(file: UploadFile = File(...)):
    """
    Upload handwritten PDF or DOCX → start chat session → return first AI question.
    """
    try:
        filename = file.filename.lower()
        file_path = os.path.join(TEMP_FOLDER, filename)
        with open(file_path, "wb") as f:
            f.write(await file.read())

        ocr_texts = []
        image_base64_list = []
        pages_processed = 0

        # Process PDF
        if filename.endswith(".pdf"):
            pages = convert_from_path(file_path, dpi=300)
            for i, page in enumerate(pages):
                img_path = os.path.join(
                    TEMP_FOLDER, f"{os.path.splitext(filename)[0]}_page_{i+1}.png"
                )
                page.save(img_path, "PNG")

                # OCR
                text = pytesseract.image_to_string(page, lang="eng").strip()
                ocr_texts.append(text)

                # Convert to base64 for AI
                with open(img_path, "rb") as img_file:
                    img_b64 = base64.b64encode(img_file.read()).decode("utf-8")
                image_base64_list.append(img_b64)

            pages_processed = len(pages)

        # Process DOCX
        elif filename.endswith(".docx"):
            from docx import Document
            doc = Document(file_path)
            full_text = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
            doc_text = "\n".join(full_text)
            ocr_texts.append(doc_text)
            pages_processed = 1

        else:
            return JSONResponse(
                {"status": "error", "message": "Unsupported file type. Upload PDF or DOCX."},
                status_code=400,
            )
     


        combined_text = "\n\n".join(ocr_texts)
        print("OACR extracted text:", combined_text[:500])
        safe_prompt = get_supervision_prompt(combined_text)

        # ----------------- Initialize LangChain Memory -----------------
        session_id = str(uuid.uuid4())
        memory = ConversationBufferMemory(return_messages=True)
        session_memories[session_id] = memory

        # Create system + user messages using LangChain schema
        system_msg = SystemMessage(
            content="You are an educational assistant that gives safe, respectful feedback to students based on their handwritten answers."
        )

        # Combine extracted text and images into a single human message
        human_content = safe_prompt
        memory.chat_memory.add_message(system_msg)
        memory.chat_memory.add_message(HumanMessage(content=human_content))

        # Call AI model
        response = client.chat.completions.create(
            model=DEPLOYMENT_NAME,
            messages=[
                {"role": "system", "content": system_msg.content},
                {"role": "user", "content": safe_prompt},
            ],
            max_tokens=800,
        )

        ai_reply = response.choices[0].message.content
        memory.chat_memory.add_message(AIMessage(content=ai_reply))

        return JSONResponse({
            "status": "success",
            "session_id": session_id,
            "pages_processed": pages_processed,
            "ai_message": ai_reply
        })

    except Exception as e:
        return JSONResponse({"status": "error", "message": str(e)}, status_code=500)


# ----------------- SEND MESSAGE -----------------
@assign_route.post("/send-message")
async def send_message(session_id: str = Form(...), student_message: str = Form(...)):
    """
    Send student's reply → get next AI message → maintain session.
    """
    try:
        if session_id not in session_memories:
            return JSONResponse({"status": "error", "message": "Invalid session_id"}, status_code=400)

        memory = session_memories[session_id]
        memory.chat_memory.add_message(HumanMessage(content=student_message))

        # Build context for AI
        messages = []
        for m in memory.chat_memory.messages:
            if isinstance(m, SystemMessage):
                messages.append({"role": "system", "content": m.content})
            elif isinstance(m, HumanMessage):
                messages.append({"role": "user", "content": m.content})
            elif isinstance(m, AIMessage):
                messages.append({"role": "assistant", "content": m.content})

        # Generate AI response
        response = client.chat.completions.create(
            model=DEPLOYMENT_NAME,
            messages=messages,
            max_tokens=800,
        )
        ai_reply = response.choices[0].message.content

        memory.chat_memory.add_message(AIMessage(content=ai_reply))

        return JSONResponse({"status": "success", "ai_message": ai_reply})

    except Exception as e:
        return JSONResponse({"status": "error", "message": str(e)}, status_code=500)


# ----------------- SESSION HISTORY -----------------
@assign_route.get("/session-history/")
async def session_history(session_id: str):
    """
    Retrieve full chat history for a session.
    """
    if session_id not in session_memories:
        return JSONResponse({"status": "error", "message": "Invalid session_id"}, status_code=400)

    memory = session_memories[session_id]
    history = [
        {"role": m.type, "content": m.content}
        for m in memory.chat_memory.messages
    ]

    return JSONResponse({"status": "success", "history": history})






# from fastapi import  UploadFile, File, Form,APIRouter
# from fastapi.responses import JSONResponse

# import os, io, base64, uuid
# from pdf2image import convert_from_path
# from PIL import Image
# import pytesseract
# from .model import get_client
# from .prompts import get_supervision_prompt


# TEMP_FOLDER = os.path.join(os.path.dirname(__file__), "pdf_images")


# # In-memory session storage
# chat_sessions = {}  # { session_id: [messages] }

# client = get_client()
# DEPLOYMENT_NAME = os.getenv("AZURE_OPENAI_DEPLOYMENT")
# assign_route = APIRouter(
#     prefix="/assignment",
#     tags=["assignment"]
# )

# @assign_route.post("/start-session")
# async def start_session(file: UploadFile = File(...)):
#     """
#     Upload handwritten PDF → start chat session → return first AI question.
#     """
#     try:
#         filename = file.filename.lower()
#         file_path = os.path.join(TEMP_FOLDER, filename)
#         with open(file_path, "wb") as f:
#             f.write(await file.read())

#         ocr_texts = []
#         image_base64_list = []
#         pages_processed = 0 

#         if filename.endswith(".pdf"):
           
#             pages = convert_from_path(file_path, dpi=300)
#             print("Pages:", len(pages))
#             print(pytesseract.image_to_string(pages[0]))
#             for i, page in enumerate(pages):
#                 img_path = os.path.join(TEMP_FOLDER, f"{os.path.splitext(filename)[0]}_page_{i+1}.png")
#                 page.save(img_path, "PNG")

#                 text = pytesseract.image_to_string(page, lang="eng").strip()
#                 ocr_texts.append(text)

#                 with open(img_path, "rb") as img_file:
#                     img_b64 = base64.b64encode(img_file.read()).decode("utf-8")
#                 image_base64_list.append(img_b64)

#         elif filename.endswith(".docx"):
#             from docx import Document
#             doc = Document(file_path)
#             full_text = []
#             for para in doc.paragraphs:
#                 if para.text.strip():
#                     full_text.append(para.text.strip())
#             doc_text = "\n".join(full_text)
#             ocr_texts.append(doc_text)
#             pages_processed = 1

#         else:
#             return JSONResponse({"status": "error", "message": "Unsupported file type. Upload PDF or DOCX."}, status_code=400)

#         combined_text = "\n\n".join(ocr_texts)
#         safe_prompt = get_supervision_prompt(combined_text)
#         print("ocr")
        

#         # Initialize session
#         session_id = str(uuid.uuid4())
#         session_messages = [
#             {"role": "system", "content": "You are an educational assistant that gives safe, respectful feedback to students based on their handwritten answers."},
#             {
#                 "role": "user",
#                 "content": [
#                     {"type": "text", "text": safe_prompt},
#                     *[
#                         {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{img_b64}"}}
#                         for img_b64 in image_base64_list
#                     ]
#                 ]
#             }
#         ]
#         chat_sessions[session_id] = session_messages

#         # Get AI first response
#         response = client.chat.completions.create(
#             model=DEPLOYMENT_NAME,
#             messages=session_messages,
#             max_tokens=800
#         )
#         ai_reply = response.choices[0].message.content
#         chat_sessions[session_id].append({"role": "assistant", "content": ai_reply})

#         return JSONResponse({
#             "status": "success",
#             "session_id": session_id,
#             "pages_processed": pages_processed,
#             "ai_message": ai_reply
#         })

#     except Exception as e:
#         return JSONResponse({"status": "error", "message": str(e)}, status_code=500)


# @assign_route.post("/send-message")
# async def send_message(session_id: str = Form(...), student_message: str = Form(...)):
#     """
#     Send student's reply → get next AI message → maintain session.
#     """
#     try:
#         if session_id not in chat_sessions:
#             return JSONResponse({"status": "error", "message": "Invalid session_id"}, status_code=400)

#         session_messages = chat_sessions[session_id]

#         # Append student's message
#         session_messages.append({"role": "user", "content": student_message})

#         # Get AI reply
#         response = client.chat.completions.create(
#             model=DEPLOYMENT_NAME,
#             messages=session_messages,
#             max_tokens=800
#         )

#         ai_reply = response.choices[0].message.content
#         session_messages.append({"role": "assistant", "content": ai_reply})

#         return JSONResponse({
#             "status": "success",
#             "ai_message": ai_reply
#         })

#     except Exception as e:
#         return JSONResponse({"status": "error", "message": str(e)}, status_code=500)

# @assign_route.get("/session-history/")
# async def session_history(session_id: str):
#     """
#     Retrieve full chat history for a session.
#     """
#     if session_id not in chat_sessions:
#         return JSONResponse({"status": "error", "message": "Invalid session_id"}, status_code=400)

#     return JSONResponse({
#         "status": "success",
#         "history": chat_sessions[session_id]
#     })
